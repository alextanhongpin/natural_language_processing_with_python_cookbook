import docx


def get_text(filename):
    doc = docx.Document(filename)
    
    # Print info.
    print('Number of paragraphs:', len(doc.paragraphs))
    
    text = []
    for idx, p in enumerate(doc.paragraphs):
        text.append(p.text)
        print(f'Paragraph {idx+1}:', doc.paragraphs[idx].text)
        print(f'Paragraph {idx+1} style:', doc.paragraphs[idx].style)
        print(f'Paragraph {idx+1} runs:')
        for i, run in enumerate(doc.paragraphs[idx].runs):
            print(f'Run {i+1}:', run.text, run.bold, run.italic, run.underline)
        print()
    
    return '\n'.join(text)